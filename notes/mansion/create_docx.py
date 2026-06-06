from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── ページ設定（A4）─────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.0)
section.right_margin  = Cm(2.0)
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)

# ── カラー定義 ────────────────────────────────
ACCENT = RGBColor(0x4A, 0x7C, 0xB4)
LIGHT  = RGBColor(0xE4, 0xEF, 0xF7)
PALE   = RGBColor(0xF4, 0xF7, 0xFB)
GRAY   = RGBColor(0xF5, 0xF5, 0xF5)
AMBER  = RGBColor(0xFF, 0xF3, 0xDC)
DARK   = RGBColor(0x22, 0x22, 0x22)
MID    = RGBColor(0x4A, 0x7C, 0xB4)
SUB    = RGBColor(0x55, 0x55, 0x55)
WARN   = RGBColor(0x8B, 0x60, 0x00)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
AMBER_B= RGBColor(0xE0, 0xA0, 0x30)

# ── ユーティリティ ────────────────────────────

def rgb_hex(rgb):
    return f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}'

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex(rgb))
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom),
                      ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), val.get('color', 'auto'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def para_border_bottom(para, color='4A7CB4', sz=12):
    """段落の下にボーダーラインを引く"""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_para_shading(para, rgb: RGBColor):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex(rgb))
    pPr.append(shd)

def add_section_title(doc, text):
    """セクション見出し：青線 + 青太字"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    para_border_bottom(p, color='4A7CB4', sz=10)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = MID
    return p

def add_sub_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = MID
    return p

def add_body(doc, text, size=12, color=DARK, italic=False, space_before=2, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.italic = italic
    return p

def add_bullet(doc, text, size=12):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = DARK
    return p

def add_note_box(doc, lines, bg=AMBER):
    """淡黄色の議題ボックス"""
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Cm(0.3)
        p.paragraph_format.right_indent = Cm(0.3)
        set_para_shading(p, bg)
        run = p.add_run(line)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x6B, 0x40, 0x00)
        if i == 0:
            run.bold = True

def make_table(doc, headers, rows, col_widths, hdr_size=12, data_size=12):
    n_cols = len(headers)
    table = doc.add_table(rows=len(rows)+1, cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # 列幅
    for i, w in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = Cm(w)

    # ヘッダ
    for ci, hd in enumerate(headers):
        cell = table.cell(0, ci)
        set_cell_bg(cell, LIGHT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(hd)
        run.bold = True
        run.font.size = Pt(hdr_size)
        run.font.color.rgb = MID

    # データ行
    for ri, row_data in enumerate(rows):
        bg = PALE if ri % 2 == 0 else RGBColor(0xFF,0xFF,0xFF)
        for ci, val in enumerate(row_data):
            cell = table.cell(ri+1, ci)
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(data_size)
            run.font.color.rgb = DARK
            if val in ('**2026年**', '**2027年**', '**2028年**') or \
               str(val).startswith('**'):
                run.bold = True
                run.font.color.rgb = MID
                set_cell_bg(cell, LIGHT)

    return table

def add_page_break(doc):
    doc.add_page_break()

def add_thin_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    para_border_bottom(p, color='D6E4F7', sz=6)

# ══════════════════════════════════════════════════
# 表紙
# ══════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.space_after  = Pt(4)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('コンフォール井の頭公園南')
run.font.size = Pt(16)
run.font.color.rgb = MID

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after  = Pt(8)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_border_bottom(p, color='4A7CB4', sz=6)
run = p.add_run('修繕委員会　第1回定例（キックオフ）')
run.font.size = Pt(26)
run.bold = True
run.font.color.rgb = DARK

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after  = Pt(4)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2026年6月7日（日）　14:00〜15:00')
run.font.size = Pt(14)
run.font.color.rgb = SUB

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(40)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('三鷹市下連雀八丁目地区公会堂')
run.font.size = Pt(13)
run.font.color.rgb = SUB

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(4)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_shading(p, PALE)
run = p.add_run('出席：小橋　／　石川　／　十文字　／　阿部　／　武藤　／　新関')
run.font.size = Pt(13)
run.font.color.rgb = DARK

# ══════════════════════════════════════════════════
# 1. 委員会の目的・アジェンダ
# ══════════════════════════════════════════════════
add_page_break(doc)

add_section_title(doc, '1．委員会の目的')

for item in [
    '● 建物・設備の状態を把握し、長期修繕計画の策定と必要に応じた修繕積立金の見直しを推進',
    '● 住民の意見を反映した計画立案と情報共有を支援',
    '● 決定は理事会・総会が行い、委員会は調査・提案・情報発信を担う',
]:
    add_body(doc, item, size=13)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(6)
set_para_shading(p, PALE)
run = p.add_run('位置づけ：理事会と連携する住民委員会　／　決定機関は理事会・総会')
run.font.size = Pt(12)
run.font.color.rgb = SUB
run.font.italic = True

add_section_title(doc, '2．本日のアジェンダ')

agenda = [
    ('1', 'メンバー紹介・役割分担'),
    ('2', '活動スコープ・スケジュールの確認'),
    ('3', 'コンサル業者選定の進め方'),
    ('4', '省エネ改修・緊急修繕の総会提案（新テーマ）'),
    ('5', '運営ルール・決定事項・次回予定'),
]
make_table(doc, ['#', '議題'], agenda,
           col_widths=[1.0, 15.0], hdr_size=12, data_size=13)

# ══════════════════════════════════════════════════
# 2. メンバー・役割分担
# ══════════════════════════════════════════════════
add_page_break(doc)

add_section_title(doc, '3．メンバー紹介・役割分担')

add_sub_heading(doc, '出席メンバー')
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(8)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_shading(p, PALE)
run = p.add_run('小橋　　石川　　十文字　　阿部　　武藤　　新関')
run.font.size = Pt(15)
run.bold = True
run.font.color.rgb = DARK

add_sub_heading(doc, '【議題】役割分担の決定')

make_table(doc,
    ['役割', '内容'],
    [
        ['委員長',   '委員会の取りまとめ・対外窓口・理事会への報告'],
        ['副委員長', '委員長補佐・議事録管理（AIツール活用）・情報共有の整理'],
    ],
    col_widths=[3.5, 13.5], hdr_size=12, data_size=13
)

doc.add_paragraph()
add_note_box(doc, [
    '追加担当について（本日議論）',
    '● 積立金評価担当（修繕積立金の適正性検証）を設けるか',
    '● その他、必要な担当・役割があるかを議論する',
])

# ══════════════════════════════════════════════════
# 3. 活動スコープ・改定スケジュール
# ══════════════════════════════════════════════════
add_page_break(doc)

add_section_title(doc, '4．活動スコープ・改定スケジュール（2026〜2028年）')

add_sub_heading(doc, '活動スコープ（3本柱）')

scope = [
    ('①', '長期修繕計画の策定（30年）',
     'コンサル選定 → 現地調査 → 計画書作成 → 2028年総会議決'),
    ('②', '短期修繕計画の立案（直近5年）',
     '優先修繕箇所の整理 → コンサル提案 → 2028年総会議決 → 工事開始'),
    ('③', '省エネ改修・緊急修繕の提案【NEW】',
     '断熱窓を起点に → 2028年総会議決 → 工事開始（補助金活用）'),
]
for num, title, detail in scope:
    bg = AMBER if num == '③' else PALE
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.0)
    set_para_shading(p, bg)
    r1 = p.add_run(f'{num}　{title}　　')
    r1.bold = True
    r1.font.size = Pt(13)
    r1.font.color.rgb = DARK
    r2 = p.add_run(detail)
    r2.font.size = Pt(12)
    r2.font.color.rgb = SUB

add_body(doc,
    '※キックオフが5月予定から6月7日に遅延。以下は1ヶ月後ろ倒しの改定版。',
    size=11, color=WARN, italic=True, space_before=8, space_after=4)

add_sub_heading(doc, '① 長期修繕計画スケジュール')
make_table(doc,
    ['時期', '内容', '主体'],
    [
        ['2026年', '', ''],
        ['6月〜',  'コンサル3社へプレゼン依頼・日程調整', '委員会'],
        ['7月',    'コンサル3社プレゼン実施（TDS・オフィスレコン・センターオフィス）', '委員会'],
        ['8月',    '業者選定・理事会決議', '委員会＋理事会'],
        ['10月',   '現地調査（委員会立会い）', '委員会＋理事会'],
        ['11月',   '調査報告・中長期修繕計画 納品・理事会報告', '理事会'],
        ['12月',   '住民向け全戸配布・修繕積立金変更案の検討開始', '委員会＋理事会'],
        ['2027年', '', ''],
        ['1〜6月', '修繕積立金変更案の審議・精査', '委員会＋理事会'],
        ['2027年総会', '中間報告（修繕計画の方向性確認）', '総会'],
        ['7〜12月', '計画の詳細化・積立金変更額の最終調整', '委員会＋理事会'],
        ['2028年', '', ''],
        ['1〜2月', '総会提案資料の最終作成・理事会承認', '委員会＋理事会'],
        ['2028年総会', '長期修繕計画・修繕積立金変更の最終議決', '総会'],
    ],
    col_widths=[2.8, 11.2, 3.0], hdr_size=12, data_size=12
)

doc.add_paragraph()
add_sub_heading(doc, '② 短期修繕計画（直近5年）・省エネ改修スケジュール')
make_table(doc,
    ['時期', '内容', '主体'],
    [
        ['2026年', '', ''],
        ['8月',    '要求仕様提示（インフラ緊急性・省エネ優先）', '委員会'],
        ['10月',   '現地調査後の優先修繕案提示', '委員会'],
        ['12月',   '短期修繕計画案の検討・補助金調査結果報告', '委員会＋理事会'],
        ['2027年', '', ''],
        ['2027年総会', '中間報告（短期修繕計画の方向性確認）', '総会'],
        ['4〜12月', '計画の詳細化・見積取得・事業者選定準備・補助金申請準備', '委員会＋理事会'],
        ['2028年', '', ''],
        ['1〜2月', '総会提案資料の最終作成・理事会承認', '委員会＋理事会'],
        ['2028年総会', '短期修繕計画・省エネ改修の議決（工事承認）', '総会'],
        ['2028年度〜', '優先修繕工事の開始（緊急度順）・断熱窓改修の実施（補助金活用）', '理事会'],
    ],
    col_widths=[2.8, 11.2, 3.0], hdr_size=12, data_size=12
)

# ══════════════════════════════════════════════════
# 4. コンサル業者選定
# ══════════════════════════════════════════════════
add_page_break(doc)

add_section_title(doc, '5．コンサル業者選定の進め方')

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(8)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_shading(p, PALE)
run = p.add_run('候補3社：　TDS　／　オフィスレコン　／　センターオフィス')
run.font.size = Pt(14)
run.bold = True
run.font.color.rgb = MID

add_sub_heading(doc, 'プレゼン実施方針')
for item in [
    '● 7月中に3社それぞれ30分のプレゼンを実施',
    '● 場所：会場を1日確保して3社連続、または管理事務室（Web参加可）',
    '● 本日：各社への連絡・日程調整・会場確保の担当を決定',
]:
    add_body(doc, item, size=13)

add_sub_heading(doc, '選定基準（3点）')
make_table(doc,
    ['#', '基準'],
    [
        ['①', '要求仕様を満たしているか'],
        ['②', '金額が妥当か（想定100〜200万円）'],
        ['③', '短期計画・省エネ改修のコンサル提案が可能か'],
    ],
    col_widths=[1.0, 16.0], hdr_size=12, data_size=13
)

add_sub_heading(doc, '依頼内容')
for item in [
    '● 劣化診断調査（第3回大規模修繕に向けた建物調査）',
    '● 30年中長期修繕計画の作成・修繕積立金の評価',
    '● 短期修繕計画（5年）の立案（追加依頼）',
    '● 省エネ改修提案（追加依頼）',
]:
    add_body(doc, item, size=13)

doc.add_paragraph()
add_note_box(doc, [
    '【議題】見積事業者選定プロセスの追加について',
    '● 管理会社任せにせず、委員会として選定基準・比較評価プロセスを持つことで透明性が高まる',
    '● 一方、現時点では計画策定が優先。コンサル選定後に整備する選択肢もある',
    '● 論点：今の段階で仕組みを検討するか、後回しにするか',
])

# ══════════════════════════════════════════════════
# 5. 省エネ改修・緊急修繕
# ══════════════════════════════════════════════════
add_page_break(doc)

add_section_title(doc, '6．新テーマ：省エネ改修・緊急修繕の2028年総会提案')

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(6)
set_para_shading(p, PALE)
run = p.add_run('長期計画策定と並行し、早期に独立した検討ラインで動く。各提案は修繕積立金への影響を確認しながら進める。')
run.font.size = Pt(12)
run.font.color.rgb = SUB
run.font.italic = True

make_table(doc,
    ['テーマ', '内容', '目標'],
    [
        ['省エネ改修',
         '断熱窓改修を優先検討（専有部対象）\n補助金・助成金（都・市・国）の活用可能性を調査\nコンサルにも提案を依頼',
         '2028年総会への提案・工事開始'],
        ['緊急修繕',
         'インフラ系（給排水管・電気設備等）の現状把握を優先\nコンサル調査と並行して早期実態把握\n緊急度が高いものは総会を待たず理事会決議で対応',
         '随時・理事会判断'],
    ],
    col_widths=[2.5, 11.0, 3.5], hdr_size=12, data_size=12
)

doc.add_paragraph()
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(2)
set_para_shading(p, PALE)
run = p.add_run('積立金評価担当が設置された場合：')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = MID

for item in [
    '● 省エネ改修：補助金差引後の実質負担と積立金残高を照合',
    '● 緊急修繕：修繕費用の積立金への影響を試算・確認し、緊急度判断の材料とする',
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    set_para_shading(p, PALE)
    run = p.add_run(item)
    run.font.size = Pt(12)
    run.font.color.rgb = SUB

# ══════════════════════════════════════════════════
# 6. 運営ルール・決定事項・次回予定
# ══════════════════════════════════════════════════
add_page_break(doc)

add_section_title(doc, '7．運営ルール')
make_table(doc,
    ['項目', '内容'],
    [
        ['打合せ頻度', '月1回程度'],
        ['コミュニケーション', 'Google Meet・チャット＋メール、必要に応じ書面・対面'],
        ['議事録', '副委員長（AIツール活用）'],
    ],
    col_widths=[4.0, 13.0], hdr_size=12, data_size=13
)

add_section_title(doc, '8．本日の決定事項')
for item in [
    '□　役割分担（委員長・副委員長・その他担当）',
    '□　コンサル3社への連絡・プレゼン日程調整・会場確保の担当',
    '□　見積事業者選定プロセスを今設けるか後回しにするか',
    '□　省エネ改修・緊急修繕の優先確認事項',
    '□　次回開催日程',
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.3)
    run = p.add_run(item)
    run.font.size = Pt(13)
    run.font.color.rgb = DARK

add_section_title(doc, '9．次回（7月）　コンサル3社プレゼン実施 → 業者選定の議論')
for item in [
    '● コンサル3社（TDS・オフィスレコン・センターオフィス）のプレゼンを実施（各30分）',
    '● 各社の提案内容・金額・対応範囲を委員会内で評価・比較',
    '● 業者選定の方向性を議論し、8月の理事会決議へ向けて準備',
]:
    add_body(doc, item, size=13)

# ══════════════════════════════════════════════════
# 参考資料（別添）
# ══════════════════════════════════════════════════
add_page_break(doc)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(4)
run = p.add_run('【参考資料】第2回定例以降の議題準備（修繕積立金の適正化）')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = SUB

add_body(doc, '※第2回定例以降の議題準備のための参考情報。本日は配布のみ。',
         size=11, color=WARN, italic=True)

add_sub_heading(doc, '積立金の「起点となる考え方」')
make_table(doc,
    ['', 'A：最低限性能維持型', 'B：資産価値維持・向上型'],
    [
        ['目的',     '安全性・基本機能の確保', '市場競争力・快適性・資産価値の持続'],
        ['修繕範囲', '安全上必要な工事のみ',   '設備更新・美観・グレードアップ含む'],
        ['積立水準', '低め',                    '高め'],
        ['リスク',   '将来の修繕不足・資産劣化', '月額負担増・合意形成が難しい'],
    ],
    col_widths=[2.5, 7.5, 7.0], hdr_size=12, data_size=12
)
add_body(doc,
    '主流はB型（資産価値維持・向上型）。国交省ガイドライン自体が「資産価値の維持・向上」を目的に明示。'
    '井の頭公園南の希少立地を考慮すると、B型を起点とした計画が推奨。',
    size=12, color=SUB, italic=True, space_before=4)

add_sub_heading(doc, '一般的な検証手法')
make_table(doc,
    ['手法', '内容', '適した場面'],
    [
        ['㎡単価方式',    '専有面積あたりの月額をガイドライン目安と比較', '現状の簡易チェック'],
        ['実積算方式',    '工事項目ごとに数量・単価・周期を積み上げ', 'コンサル選定後の精緻な見直し'],
        ['均等積立方式の検証', '計画期間を通じて月額が適正か確認', '段階増額からの見直し時'],
    ],
    col_widths=[3.5, 8.5, 5.0], hdr_size=12, data_size=12
)
add_body(doc,
    '国交省ガイドライン（令和6年改定）の目安単価：中規模マンションで月額175〜202円/㎡',
    size=12, color=WARN, space_before=4)

add_sub_heading(doc, '三鷹市 マンション管理計画認定制度')
add_body(doc, '2024年4月より三鷹市が導入。管理組合の管理計画が国交省基準を満たす場合、市長の認定を受けられる制度。', size=12)
make_table(doc,
    ['項目', '内容'],
    [
        ['メリット①', '適正管理マンションとして市場評価が向上（売却・賃貸に有利）'],
        ['メリット②', '住宅金融支援機構「フラット35」等の金利優遇'],
        ['メリット③', '大規模修繕実施時に固定資産税（建物部分）が減額される場合あり'],
        ['申請手続き', 'マンション管理センターで事前確認 → 三鷹市へ申請（手数料無料）'],
        ['有効期間',   '5年間（更新申請が必要）'],
    ],
    col_widths=[3.0, 14.0], hdr_size=12, data_size=12
)
add_body(doc,
    '→ 長期修繕計画が整備されれば認定申請の要件を満たしやすくなる。今回の計画策定と並行して検討する価値あり。',
    size=12, color=MID, space_before=4)

# ════════════════════════════════════════════════
out = '/home/user/Ryuichi_Ishikawa_Private/notes/mansion/修繕委員会_第1回定例_キックオフ.docx'
doc.save(out)
print('saved:', out)
